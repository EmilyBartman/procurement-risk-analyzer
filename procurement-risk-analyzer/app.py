# -*- coding: utf-8 -*-
"""app.ipynb

Automatically generated by Colab.

Original file is located at
    https://colab.research.google.com/drive/1p7Ef13LULpVGcvXTjI7MQHdrFr51zMNb
"""

# === STREAMLIT DEPLOYMENT VERSION ===

# STEP 1: Import Required Libraries
import os
import io
import glob
from pathlib import Path
import pandas as pd
import streamlit as st
from dotenv import load_dotenv
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain_unstructured import UnstructuredLoader
from langchain_community.vectorstores import FAISS
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain.schema import Document as LCDocument
from PyPDF2 import PdfReader
from docx import Document
import warnings
import shutil
import streamlit as st
import streamlit.components.v1 as components


warnings.filterwarnings("ignore", category=DeprecationWarning)
warnings.filterwarnings("ignore", category=pd.errors.ParserWarning)

# STEP 2: Load Environment Variables
load_dotenv()
IFI_API_KEY = os.getenv("IFI_API_KEY")  # <-- ADD YOUR API KEY to .streamlit/secrets.toml or env vars

# STEP 3: Ensure necessary folders exist before file operations
for folder in ['historical_documents', 'risks_document', 'target_document', 'outputs']:
    Path(folder).mkdir(parents=True, exist_ok=True)

# STEP 4: Define the RAG Risk Analysis Class
class RAGProcurementRisksAnalysis:
    def __init__(self, api_key, query, historical_documents_folder_path, risks_document_folder_path, target_document_folder_path, risk_analysis_output_path):
        self.api_key = api_key
        self.query = query
        self.historical_documents = self.load_documents(historical_documents_folder_path)
        self.risks_document = self.load_documents(risks_document_folder_path)
        self.target_document = self.load_documents(target_document_folder_path)
        self.risk_analysis_output_path = risk_analysis_output_path

    def load_documents(self, folder_path):
        all_documents = []
        supported_exts = ["csv", "pdf", "docx"]
        for ext in supported_exts:
            files = glob.glob(f"{folder_path}/*.{ext}")
            for file_path in files:
                try:
                    if file_path.endswith(".csv"):
                        try:
                            with open(file_path, "r", encoding="utf-8") as f:
                                content = f.read()
                        except UnicodeDecodeError:
                            with open(file_path, "r", encoding="latin1") as f:
                                content = f.read()
                        doc = LCDocument(page_content=content)
                        all_documents.append(doc)

                    else:
                        loader = UnstructuredLoader(file_path=file_path)
                        documents = loader.load()
                        all_documents.extend(documents)
                except Exception as e:
                    print(f"⚠️ Could not load {file_path}: {e}")
        print(f"📄 Loaded {len(all_documents)} docs from {folder_path}")
        return all_documents

    def create_embeddings(self):
        embeddings = OpenAIEmbeddings(openai_api_key=self.api_key)
        vector_store = FAISS.from_documents(self.historical_documents, embeddings)
        return vector_store, embeddings

    def semantic_search(self):
        vector_store, embeddings = self.create_embeddings()
        query_embedding = embeddings.embed_query(self.query)
        risks_document_embedding = embeddings.embed_query(self.risks_document[0].page_content)
        target_document_embedding = embeddings.embed_query(self.target_document[0].page_content)
    
        retrieved_by_query = vector_store.similarity_search_by_vector(query_embedding, k=3)
        retrieved_by_risks = vector_store.similarity_search_by_vector(risks_document_embedding, k=3)
        retrieved_by_target = vector_store.similarity_search_by_vector(target_document_embedding, k=3)
    
        retrieved_documents = list({doc.page_content: doc for doc in retrieved_by_query + retrieved_by_target + retrieved_by_risks}.values())

        if not retrieved_documents:
            print("⚠️ No documents retrieved during semantic search!")
    
        print(f"🔍 Retrieved {len(retrieved_documents)} relevant docs for semantic search.")
    
        return "\n\n".join([f"Document {i + 1}: {doc.page_content}" for i, doc in enumerate(retrieved_documents)])

    def save_risk_analysis_to_file(self, risk_analysis):
        file_path = f"{self.risk_analysis_output_path}/risk_analysis.txt"
        os.makedirs(self.risk_analysis_output_path, exist_ok=True)
        with open(file_path, "w") as file:
            file.write(risk_analysis)
    
    def generate_risks_analysis_rag(self):
        llm = ChatOpenAI(model="gpt-4o", temperature=0.5, openai_api_key=self.api_key)
    
        retrieved_docs_str = self.semantic_search()
        risks_content = self.risks_document[0].page_content
        target_content = self.target_document[0].page_content
    
        # Add fallback for empty inputs
        if not retrieved_docs_str.strip():
            retrieved_docs_str = "No relevant documents were retrieved. Please proceed with only risks and target documents."
    
        if not risks_content.strip():
            st.error("❌ The risks document is empty. Please upload a valid file.")
            return "Error: Risks document is empty."
    
        if not target_content.strip():
            st.error("❌ The target document is empty. Please upload a valid file.")
            return "Error: Target document is empty."
    
        # Debug logs
        print("----- Prompt Preview -----")
        print("Query:", self.query)
        print("--- Retrieved Docs ---")
        print(retrieved_docs_str[:500])
        print("--- Risks Document ---")
        print(risks_content[:500])
        print("--- Target Document ---")
        print(target_content[:500])

    
        prompt_template = PromptTemplate(
            input_variables=["retrieved_docs_str", "risks_document_content", "target_document_content"],
            template='''You are a procurement risk assessment AI. Evaluate the risks associated with the target document
    based on the retrieved knowledge and the risks detailed in the risks document.
    
    ### Target Document:
    {target_document_content}
    
    ### Risks Document:
    {risks_document_content}
    
    ### Retrieved Risk-Related Documents:
    {retrieved_docs_str}
    
    ### Task:
    Analyze the target document and classify risks into the categories detailed in the risks document.
    
    Output the risk labels and a short explanation for each.
    
    Risk Assessment:
    
    Based on the risks document summarize a mitigation plan.
    
    Mitigation Plan:'''
        )
    
        chain = LLMChain(llm=llm, prompt=prompt_template)
        risk_analysis = chain.run({
            "retrieved_docs_str": retrieved_docs_str,
            "risks_document_content": risks_content,
            "target_document_content": target_content
        })
    
        self.save_risk_analysis_to_file(risk_analysis)
        return risk_analysis


# STEP 5: Preview Function
def preview_file(file, file_type, name="Uploaded file"):
    st.subheader(f"Preview: {name}")
    if file_type == "csv":
        df = pd.read_csv(file)
        st.dataframe(df.head())
    elif file_type == "pdf":
        reader = PdfReader(file)
        text = "\n".join([page.extract_text() for page in reader.pages[:2] if page.extract_text()])
        st.text_area("PDF Preview", text[:2000], height=200)
    elif file_type == "docx":
        doc = Document(file)
        text = "\n".join([p.text for p in doc.paragraphs])
        st.text_area("DOCX Preview", text[:2000], height=200)

# STEP 6: Streamlit UI Setup
st.set_page_config(page_title="Procurement Risk Analyzer", layout="centered")

st.title("📄 Procurement Risk Analyzer")

st.sidebar.title("ℹ️ About")
st.sidebar.info('''
This tool uses LLM-based Retrieval-Augmented Generation (RAG) to assess risks in procurement documents.

- Built with LangChain + Streamlit
- Supports CSV, PDF, and DOCX
- Preview documents before analysis
- Securely runs in your environment
''')

with st.expander("📁 Download Example Templates", expanded=True):
    col1, col2, col3 = st.columns(3)
    with col1:
        EXAMPLES_PATH = Path(__file__).resolve().parent.parent / "example_files"
        with open(EXAMPLES_PATH / "dataset1.csv", "rb") as f:
            st.download_button("⬇️ History Doc.csv", f, file_name="dataset1.csv", help="Historical doc example")
    with col2:
        with open(EXAMPLES_PATH / "risks.csv", "rb") as f:
            st.download_button("⬇️ Risks Doc.csv", f, file_name="risks.csv", help="Risk types to reference")
    with col3:
        with open(EXAMPLES_PATH / "dataset_no_risks.csv", "rb") as f:
            st.download_button("⬇️ Target Doc.csv", f, file_name="dataset_no_risks.csv", help="Target doc example")

st.markdown("### Step 1: Upload Your Files")
query = st.text_input("What do you want to know?", "What are the risks associated with this procurement document?")

historical_files = st.file_uploader(
    "Upload historical documents (.csv, .pdf, .docx)",
    accept_multiple_files=True,
    type=["csv", "pdf", "docx"],
    help="📚 Historical documents are previous procurement records that help the model understand patterns and context.\n\nExample: History Doc.csv"
)

risks_file = st.file_uploader(
    "Upload a risks document (.csv, .pdf, .docx)",
    type=["csv", "pdf", "docx"],
    help="⚠️ The risks document defines risk types (e.g., Schedule Risk, Cost Risk) and their descriptions for assessment guidance.\n\nExample: Risks Doc.csv"
)

target_file = st.file_uploader(
    "Upload a target procurement document (.csv, .pdf, .docx)",
    type=["csv", "pdf", "docx"],
    help="🎯 The target document is the procurement form or data you'd like to analyze. It should contain project variables, dates, and dependencies.\n\nExample: Target Doc.csv"
)

historical_file_bytes = []
if historical_files:
    for f in historical_files:
        bytes_data = f.getvalue()
        file_ext = f.name.split(".")[-1]
        st.text(f"🧪 Uploaded historical file: {f.name}, size: {len(bytes_data)} bytes")
        preview_file(io.BytesIO(bytes_data), file_ext, name=f.name)
        historical_file_bytes.append((f.name, bytes_data))

if risks_file:
    risks_bytes = risks_file.getvalue()
    file_ext = risks_file.name.split(".")[-1]
    st.text(f"🧪 Uploaded risks file: {risks_file.name}, size: {len(risks_bytes)} bytes")
    preview_file(io.BytesIO(bytes_data), file_ext, name=f.name)


if target_file:
    target_bytes = target_file.getvalue()
    file_ext = target_file.name.split(".")[-1]
    st.text(f"🧪 Uploaded target file: {target_file.name}, size: {len(target_bytes)} bytes")
    preview_file(io.BytesIO(bytes_data), file_ext, name=f.name)



if st.button("Run Analysis"):
    if not IFI_API_KEY:
        st.error("Missing API key!")
    elif not historical_files or not risks_file or not target_file:
        st.warning("Please upload all required files.")
    else:
        with st.spinner("Processing files and analyzing..."):
            base_dir = Path(".")

            # Save historical files
            for fname, fbytes in historical_file_bytes:
                with open(base_dir / "historical_documents" / fname, "wb") as out:
                    out.write(fbytes)

            # Save risks file
            risks_path = base_dir / "risks_document" / risks_file.name
            with open(risks_path, "wb") as out:
                out.write(risks_bytes)

            # Save target file
            target_path = base_dir / "target_document" / target_file.name
            with open(target_path, "wb") as out:
                out.write(target_bytes)

            # Run RAG analysis
            rag = RAGProcurementRisksAnalysis(
                api_key=IFI_API_KEY,
                query=query,
                historical_documents_folder_path=base_dir / "historical_documents",
                risks_document_folder_path=base_dir / "risks_document",
                target_document_folder_path=base_dir / "target_document",
                risk_analysis_output_path=base_dir / "outputs"
            )

            st.text(f"📄 Loaded {len(rag.risks_document)} risks doc(s)")
            if rag.risks_document:
                st.text(f"🔎 Risks doc preview:\n{rag.risks_document[0].page_content[:300]}")

            if not rag.historical_documents:
                st.error("❌ Could not load any content from historical documents.")
            elif not rag.risks_document:
                st.error("❌ Could not load any content from the risks document.")
            elif not rag.target_document:
                st.error("❌ Could not load any content from the target document.")
            else:
                result = rag.generate_risks_analysis_rag()
                st.success("✅ Analysis complete!")

                st.download_button("📥 Download Result", result, file_name="risk_analysis.txt")

                if "Mitigation Plan:" in result:
                    risk_section, mitigation_section = result.split("Mitigation Plan:", 1)
                else:
                    risk_section = result
                    mitigation_section = ""

                with st.expander("📋 Risk Assessment", expanded=True):
                    st.markdown(risk_section.strip(), unsafe_allow_html=True)

                if mitigation_section.strip():
                    with st.expander("🛡️ Mitigation Plan", expanded=True):
                        st.markdown(mitigation_section.strip(), unsafe_allow_html=True)
